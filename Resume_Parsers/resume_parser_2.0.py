import os
import re
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import argparse  # Import argparse

# --- (All the functions from the previous response go here) ---
# TECHNICAL_SKILLS_LIST, SKILL_SECTION_KEYWORDS, GENERAL_SECTION_DELIMITERS
# compile_regex_patterns, SKILL_SECTION_HEADER_REGEX, DELIMITER_HEADER_REGEX
# COMPILED_SKILL_PATTERNS, SINGLE_WORD_SKILLS_SET
# extract_text_from_pdf, extract_text_from_docx, extract_text_from_txt
# get_text_from_resume, preprocess_text_for_skill_matching
# extract_targeted_sections, find_skills_in_chunk
# parse_resume_for_skills
# --- (Ensure NLTK resources are available - same as before) ---
try:
    stopwords.words('english')
    word_tokenize("test")
except LookupError:
    print("NLTK resources not found. Downloading...")
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    print("NLTK resources downloaded.")

# !!! CRITICAL: EXPAND AND REFINE THIS LIST EXTENSIVELY !!!
TECHNICAL_SKILLS_LIST = [
    "python", "java", "c++", "c#", "javascript", "typescript", "php", "ruby", "swift", "kotlin", "go", "rust",
    "sql", "mysql", "postgresql", "mongodb", "nosql", "sqlite", "oracle database", "sql server",
    "react", "angular", "vue.js", "vue", "next.js", "gatsby",
    "django", "flask", "spring", "spring boot", ".net", "dotnet", "asp.net", "laravel", "ruby on rails",
    "node.js", "express.js",
    "html", "html5", "css", "css3", "sass", "less", "bootstrap", "tailwind css",
    "aws", "azure", "google cloud platform", "gcp", "docker", "kubernetes", "k8s", "terraform", "ansible", "jenkins",
    "linux", "unix", "bash scripting", "powershell",
    "git", "github", "gitlab", "bitbucket", "jira", "confluence",
    "rest", "restful apis", "graphql", "soap", "microservices", "api design",
    "machine learning", "deep learning", "natural language processing", "nlp", "computer vision", "data science",
    "pandas", "numpy", "scipy", "scikit-learn", "sklearn", "tensorflow", "keras", "pytorch",
    "big data", "hadoop", "spark", "apache spark", "kafka", "apache kafka",
    "data analysis", "data visualization", "tableau", "power bi",
    "cybersecurity", "penetration testing", "ethical hacking", "encryption", "network security",
    "agile", "scrum", "kanban", "devops", "ci/cd",
    "object-oriented programming", "oop", "functional programming",
    "selenium", "playwright", "cypress", "unit testing", "integration testing",
    "r programming", "matlab", "objective-c", "perl", "scala"
]
SKILL_SECTION_KEYWORDS = [
    "skills", "technical skills", "technical proficiencies", "core competencies",
    "key skills", "technologies", "tools & technologies", "tools and technologies",
    "programming languages", "software proficiency", "technical expertise",
    "expertise", "proficient in", "familiar with"
]
GENERAL_SECTION_DELIMITERS = [
    "experience", "work experience", "professional experience", "employment history",
    "education", "academic qualifications", "qualifications", "academic background",
    "projects", "personal projects", "portfolio",
    "summary", "objective", "profile", "about me", "professional summary",
    "awards", "honors", "recognitions",
    "publications", "research",
    "certifications", "licenses", "courses",
    "references", "volunteer experience", "extracurricular activities", "interests", "hobbies"
]


def compile_regex_patterns(keywords):
    return re.compile(
        r"^\s*(" + "|".join(re.escape(k) for k in keywords) + r")\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE
    )


SKILL_SECTION_HEADER_REGEX = compile_regex_patterns(SKILL_SECTION_KEYWORDS)
DELIMITER_HEADER_REGEX = compile_regex_patterns(GENERAL_SECTION_DELIMITERS)

COMPILED_SKILL_PATTERNS = []
SINGLE_WORD_SKILLS_SET = set()
for skill in TECHNICAL_SKILLS_LIST:
    skill_lower = skill.lower()
    if " " in skill_lower or "." in skill_lower or "#" in skill_lower or "+" in skill_lower or "-" in skill_lower:  # Added hyphen
        COMPILED_SKILL_PATTERNS.append(re.compile(
            r"\b" + re.escape(skill_lower) + r"\b", re.IGNORECASE))
    else:
        SINGLE_WORD_SKILLS_SET.add(skill_lower)


def extract_text_from_pdf(file_path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except ImportError:
        print("PyPDF2 is not installed. Please install it: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except ImportError:
        print("python-docx is not installed. Please install it: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return ""


def extract_text_from_txt(file_path):
    try:
        # Added errors='ignore' for robustness
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return ""


def get_text_from_resume(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return None
    if not os.path.isfile(file_path):  # Check if it's a file
        print(f"Error: {file_path} is not a file.")
        return None

    _, extension = os.path.splitext(file_path.lower())
    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension == ".docx":
        return extract_text_from_docx(file_path)
    elif extension == ".txt":
        return extract_text_from_txt(file_path)
    else:
        print(
            f"Unsupported file format: {extension}. Please provide a .pdf, .docx, or .txt file.")
        return None


stop_words = set(stopwords.words('english'))


def preprocess_text_for_skill_matching(text):
    if not text:
        return ""
    text_lower = text.lower()
    return text_lower


def extract_targeted_sections(full_text):
    skill_sections_content = []
    if not full_text:
        return skill_sections_content  # Guard clause
    lines = full_text.splitlines()
    in_skill_section = False
    current_section_lines = []
    MAX_LINES_IN_SKILL_SECTION_WITHOUT_DELIMITER = 30  # Adjusted
    CONSECUTIVE_BLANK_LINES_TO_END_SECTION = 2
    consecutive_blanks = 0

    for line_num, line in enumerate(lines):
        line_stripped = line.strip()
        is_skill_header = bool(SKILL_SECTION_HEADER_REGEX.match(line_stripped))
        is_delimiter_header = bool(DELIMITER_HEADER_REGEX.match(
            line_stripped)) and not is_skill_header

        if is_skill_header:
            if in_skill_section and current_section_lines:
                skill_sections_content.append("\n".join(current_section_lines))
            current_section_lines = []
            in_skill_section = True
            consecutive_blanks = 0
            continue

        if in_skill_section:
            if is_delimiter_header:
                if current_section_lines:
                    skill_sections_content.append(
                        "\n".join(current_section_lines))
                current_section_lines = []
                in_skill_section = False
                consecutive_blanks = 0
                continue

            if not line_stripped:
                consecutive_blanks += 1
                # Only end section due to blanks if there's actual content collected
                if consecutive_blanks >= CONSECUTIVE_BLANK_LINES_TO_END_SECTION and current_section_lines:
                    skill_sections_content.append(
                        "\n".join(current_section_lines))
                    current_section_lines = []
                    in_skill_section = False
                    consecutive_blanks = 0
                    # Don't continue here, as this blank line might be followed by a delimiter
            else:
                consecutive_blanks = 0
                current_section_lines.append(line)  # Add original line

            if len(current_section_lines) > MAX_LINES_IN_SKILL_SECTION_WITHOUT_DELIMITER:
                if current_section_lines:
                    skill_sections_content.append(
                        "\n".join(current_section_lines))
                current_section_lines = []
                in_skill_section = False
                consecutive_blanks = 0

    if in_skill_section and current_section_lines:
        skill_sections_content.append("\n".join(current_section_lines))
    return skill_sections_content


def find_skills_in_chunk(text_chunk):
    if not text_chunk:
        return set()
    processed_chunk = preprocess_text_for_skill_matching(text_chunk)
    found_skills_in_chunk = set()

    # Regex for multi-word or special character skills
    for skill_pattern in COMPILED_SKILL_PATTERNS:
        # Find the original skill name based on the pattern's source
        # This assumes the pattern was created directly from re.escape(skill.lower())
        # A more robust mapping might be needed if patterns are complex
        original_skill_name = next((s for s in TECHNICAL_SKILLS_LIST if re.escape(
            s.lower()) in skill_pattern.pattern), None)
        if original_skill_name and skill_pattern.search(processed_chunk):
            found_skills_in_chunk.add(original_skill_name)

    # Single-word skills
    for single_skill_lower in SINGLE_WORD_SKILLS_SET:
        pattern = r"\b" + re.escape(single_skill_lower) + r"\b"
        if re.search(pattern, processed_chunk):
            original_skill_name = next(
                (s for s in TECHNICAL_SKILLS_LIST if s.lower() == single_skill_lower), None)
            if original_skill_name:
                found_skills_in_chunk.add(original_skill_name)
    return found_skills_in_chunk


def parse_resume_for_skills(resume_file_path):
    print(f"Processing resume: {resume_file_path}\n")
    full_resume_text = get_text_from_resume(resume_file_path)

    if not full_resume_text or not full_resume_text.strip():
        print(
            f"Could not extract text or text is empty from {resume_file_path}")
        return []

    all_extracted_skills = set()
    processed_full_text_effectively = False

    skill_section_texts = extract_targeted_sections(full_resume_text)
    if skill_section_texts:
        print(f"Found {len(skill_section_texts)} potential skill section(s).")
        for i, section_text in enumerate(skill_section_texts):
            print(
                f"--- Analyzing content of identified skill section {i+1} ---")
            skills_from_section = find_skills_in_chunk(section_text)
            if skills_from_section:
                print(
                    f"Skills found in section {i+1}: {', '.join(sorted(list(skills_from_section)))}\n")
                all_extracted_skills.update(skills_from_section)
                processed_full_text_effectively = True  # Mark that we got skills from sections
            else:
                print(
                    f"No specific skills from the predefined list found in section {i+1}.\n")
    else:
        print("No dedicated skill sections found using keywords.")

    # If no skills found in specific sections, or as a supplementary pass always (optional)
    if not processed_full_text_effectively:  # or you can remove this condition to always do a full scan
        if not skill_section_texts:  # Only print this if no sections were found at all
            print("\n--- Analyzing the full resume text as no specific skill sections were identified or they yielded no results ---")
        else:  # Sections were found but no skills in them
            print("\n--- Performing a supplementary scan of the full resume text as no skills were found in identified sections ---")

        skills_from_full_scan = find_skills_in_chunk(full_resume_text)
        if skills_from_full_scan:
            # Show only new skills if sections were processed
            newly_found = skills_from_full_scan - all_extracted_skills
            if newly_found:
                print(
                    f"Additional skills found in full scan: {', '.join(sorted(list(newly_found)))}\n")
                all_extracted_skills.update(skills_from_full_scan)
            elif not all_extracted_skills:  # If no skills were found before, print all from full scan
                print(
                    f"Skills found in full scan: {', '.join(sorted(list(skills_from_full_scan)))}\n")
                all_extracted_skills.update(skills_from_full_scan)

    if not all_extracted_skills:
        print("No technical skills from the predefined list were found in the resume.")
    return sorted(list(all_extracted_skills))


def run_dummy_tests():
    """Creates and tests with dummy resume files."""
    print("\n--- Running Dummy Tests ---")
    dummy_resume_text_1 = """
    John Doe
    Technical Skills:
    - Programming Languages: Python, JavaScript, Java, C#
    - Web Technologies: React, Node.js, HTML5, CSS3, RESTful APIs
    - Databases: MongoDB, PostgreSQL, MySQL. Also worked with SQL Server.
    - Cloud: AWS (EC2, S3, Lambda), Docker
    - Tools: Git, Jenkins, Jira. Familiar with C++.
    Experience
    Developed using python scripting.
    """
    dummy_resume_text_2 = """
    Jane Smith
    Profile
    My work involves extensive use of SQL for data extraction and Python (Pandas, NumPy) for data manipulation and analysis.
    I also have experience creating visualizations with Tableau. Experienced with git for version control.
    Used machine learning models (scikit-learn) for predictive analytics.
    """
    dummy_resume_text_docx = """Alice Wonderland
    Core Competencies
    Languages: Ruby, Go, Swift
    Frameworks: Ruby on Rails
    Cloud: Google Cloud Platform (GCP), Kubernetes (k8s)
    Other: CI/CD, DevOps practices, .NET experience.
    """
    dummy_files_dir = "dummy_resumes_for_testing"  # Changed name to avoid conflict
    if not os.path.exists(dummy_files_dir):
        os.makedirs(dummy_files_dir)

    resume_path_1 = os.path.join(dummy_files_dir, "JobResume2.txt")
    resume_path_2 = os.path.join(dummy_files_dir, "jane_smith_resume.txt")
    resume_path_docx = os.path.join(
        dummy_files_dir, "alice_wonderland_resume.docx")

    with open(resume_path_1, "w", encoding="utf-8") as f:
        f.write(dummy_resume_text_1)
    with open(resume_path_2, "w", encoding="utf-8") as f:
        f.write(dummy_resume_text_2)
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph(dummy_resume_text_docx)
        doc.save(resume_path_docx)
        test_files = [resume_path_1, resume_path_2, resume_path_docx]
    except ImportError:
        print("python-docx not installed. Skipping DOCX dummy file creation and test.")
        test_files = [resume_path_1, resume_path_2]
    except Exception as e:
        print(f"Error creating dummy DOCX: {e}")
        test_files = [resume_path_1, resume_path_2]

    for resume_file in test_files:
        if os.path.exists(resume_file):
            extracted_skills = parse_resume_for_skills(resume_file)
            print(
                f"\n>>> Final Extracted Skills for {os.path.basename(resume_file)}: {extracted_skills}")
            print("-" * 50)


# --- Main Execution Block ---
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Parse a resume file to extract technical skills.")
    parser.add_argument("resume_file", nargs='?',
                        help="Path to the resume file (PDF, DOCX, TXT).")
    parser.add_argument("--test", action="store_true",
                        help="Run with built-in dummy resume files.")

    args = parser.parse_args()

    if args.test:
        run_dummy_tests()
    elif args.resume_file:
        if os.path.exists(args.resume_file):
            extracted_skills = parse_resume_for_skills(args.resume_file)
            print(
                f"\n>>> Final Extracted Skills for {os.path.basename(args.resume_file)}: {extracted_skills}")
        else:
            print(f"Error: The file '{args.resume_file}' does not exist.")
            print("Please provide a valid file path.")
    else:
        # No file provided and --test not used, print help.
        parser.print_help()
        print("\nOr run with --test to use dummy files.")
        print("Example: python your_script_name.py path/to/your_resume.pdf")
        print("Example: python your_script_name.py --test")
